#!/usr/bin/env python3
"""Build the active double-anonymous JEB package from V5/V3 Markdown sources.

This reuses the rendering/layout helpers from the frozen legacy builder while
routing the active V5 text and embedding the five main figures near first mention.
Older V3/V4 builders remain audit snapshots.
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph

import build_chapter2_jeb_docx_v1 as legacy

ROOT = Path(__file__).resolve().parents[1]
CH = ROOT / "docs" / "chapter2"
OUT = CH / "submission_package_v5"

FIGURE_SPECS = [
    {
        "number": 1,
        "marker": "(Fig. 1",
        "filename": "figure1_evolutionary_depth.png",
        "caption": "Figure 1. Capitulum modules have unequal evolutionary depth. Panel a shows authority-backed trait-state coverage, panel b the minimum unordered change burden across the admitted topology ensemble, and panel c the bootstrap-median relative lineage-depth envelopes. Relative lineage-depth is topology-only and is not calendar time or an evolutionary rate.",
        "alt": "Three-panel figure comparing orientation, phyllary posture, and stickiness. Trait coverage is 20, 10, and 13 concepts; minimum changes are 4 to 6, 3, and 5; relative lineage-depth envelopes are 0.795 to 0.994, 0.695 to 1.000, and 0.937 to 0.954.",
    },
    {
        "number": 2,
        "marker": "(Fig. 2",
        "filename": "figure2_shared_history_boundary.png",
        "caption": "Figure 2. One present capitulum does not imply one transition history. Named-edge forcing differs among traits, and none of the three trait pairs passes the prespecified robust shared-transition-localization rule. This constrains a simple synchronized-history model but does not establish complete trait independence.",
        "alt": "Three-panel figure showing selected forced-edge bootstrap fractions, branch-aware and equal-branch transition-localization correlations, and a decision box stating that zero of three trait pairs passes the robust shared-localization rule.",
    },
    {
        "number": 3,
        "marker": "(Fig. 3",
        "filename": "figure3_orientation_state_trajectory.png",
        "caption": "Figure 3. Present hydric correspondence does not identify the historical origin of capitulum orientation. Current BIO12, BIO15, and BIO1 correspondences are shown separately; the candidate core-Nipponocirsium change is bounded by a cross-study chronology envelope; and 376 chronology-by-palaeolocation scenarios span both positive and negative state-trajectory cosine values, leaving the origin trajectory unresolved.",
        "alt": "Four-panel orientation figure. Present environmental effects point to a hydric domain, but the public chronology spans 94 admissible age pairs and four regional scenarios. Historical state-trajectory cosine values range from negative to positive, with q05 minus 0.799, median minus 0.065, and q95 plus 0.609, so the origin trajectory is unresolved.",
    },
    {
        "number": 4,
        "marker": "(Fig. 4",
        "filename": "figure4_colour_rsds_natural_experiment.png",
        "caption": "Figure 4. Repeated white phenotypes do not imply one current radiation context. Both dated sister systems show lower chroma and higher lightness in the white lineage, but current RSDS is higher in the Arenicola white lineage and lower in the Taiwan white lineage. The pair-level Azami-direction concordance is therefore one of two, while the pooled within-taxon chroma-RSDS slope remains negative as a secondary scale diagnostic.",
        "alt": "Four-panel colour figure comparing Arenicola and Taiwan sister systems. White lineages have lower chroma in both systems. RSDS white-minus-coloured is positive in Arenicola but negative in Taiwan, including after locality aggregation; pair-level concordance is one of two.",
    },
    {
        "number": 5,
        "marker": "(Fig. 5",
        "filename": "figure5_whole_capitulum_synthesis.png",
        "caption": "Figure 5. Partial coordinated remodelling is nested within a modular historical mosaic. Circularity and solidity are higher and visible floret fraction lower in both white-lineage comparisons, whereas finer geometry is heterogeneous or low-information. Together with the zero-of-three shared-history result and partial present integration, this supports an intermediate model between complete independence and one synchronized universal syndrome.",
        "alt": "Four-panel synthesis figure. Three coarse head traits shift in the same direction in both white sister systems, fine geometry is heterogeneous, and the final process model places partial coordinated remodelling between rejected extremes of complete independence and one universal synchronized syndrome.",
    },
]


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser()
    p.add_argument("--output-dir", type=Path, default=OUT)
    p.add_argument("--figure-dir", type=Path, default=CH / "figures_v5")
    return p.parse_args()


def blank_identifying_metadata(doc: Document) -> None:
    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.comments = ""
    props.category = ""
    props.subject = ""
    props.keywords = ""


def paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def add_figure_after(paragraph: Paragraph, image_path: Path, caption: str, alt: str) -> None:
    if not image_path.exists() or image_path.stat().st_size < 3000:
        raise RuntimeError(f"Missing or unexpectedly small figure: {image_path}")

    image_par = paragraph_after(paragraph)
    image_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = image_par.add_run().add_picture(str(image_path), width=Inches(6.35))
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", caption.split(".", 1)[0])

    caption_par = paragraph_after(image_par)
    try:
        caption_par.style = "Caption"
    except KeyError:
        pass
    caption_par.add_run(caption)


def embed_main_figures(doc: Document, figure_dir: Path) -> None:
    original_paragraphs = list(doc.paragraphs)
    used = set()
    for spec in FIGURE_SPECS:
        matches = [p for p in original_paragraphs if spec["marker"] in p.text]
        if not matches:
            raise RuntimeError(f"No first-mention paragraph found for Figure {spec['number']}")
        target = matches[0]
        if spec["number"] in used:
            raise RuntimeError(f"Figure inserted twice: {spec['number']}")
        add_figure_after(
            target,
            figure_dir / spec["filename"],
            spec["caption"],
            spec["alt"],
        )
        used.add(spec["number"])
    if used != {1, 2, 3, 4, 5}:
        raise RuntimeError(f"Incomplete figure insertion: {sorted(used)}")


def build_main(output_dir: Path, figure_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="", line_numbers=True)
    blank_identifying_metadata(doc)
    legacy.render_markdown(
        doc,
        CH / "MANUSCRIPT_JEB_V5.md",
        skip_metadata_prefixes=(
            "**Target journal:**",
            "**Manuscript status:**",
            "**Running title:**",
            "**Word-limit contract:**",
        ),
        stop_heading="Submission completion gates",
    )
    embed_main_figures(doc, figure_dir)
    path = output_dir / "Chapter2_JEB_Anonymous_Manuscript_V5.docx"
    legacy.save_document(doc, path)
    return path


def build_title_page(output_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="", line_numbers=False)
    legacy.render_markdown(doc, CH / "JEB_TITLE_PAGE_TEMPLATE_V2.md")
    path = output_dir / "Chapter2_JEB_Title_Page_TEMPLATE_V2.docx"
    legacy.save_document(doc, path)
    return path


def build_supporting(output_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="", line_numbers=True)
    legacy.render_markdown(doc, CH / "JEB_SUPPORTING_INFORMATION_V3.md")
    path = output_dir / "Chapter2_JEB_Supporting_Information_V3.docx"
    legacy.save_document(doc, path)
    return path


def build_cover_letter(output_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="", line_numbers=False)
    legacy.render_markdown(doc, CH / "JEB_COVER_LETTER_TEMPLATE_V2.md")
    path = output_dir / "Chapter2_JEB_Cover_Letter_TEMPLATE_V2.docx"
    legacy.save_document(doc, path)
    return path


def display_path(path: Path) -> str:
    try:
        return str(path.relative_to(ROOT))
    except ValueError:
        return str(path)


def main() -> int:
    args = parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    outputs = [
        build_main(args.output_dir, args.figure_dir),
        build_title_page(args.output_dir),
        build_supporting(args.output_dir),
        build_cover_letter(args.output_dir),
    ]
    for path in outputs:
        if not path.exists() or path.stat().st_size < 1000:
            raise RuntimeError(f"Document build failed or unexpectedly small: {path}")
        print(display_path(path), path.stat().st_size)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
