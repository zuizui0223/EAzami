#!/usr/bin/env python3
"""Build the active JEB V7 DOCX submission package.

Legacy builders remain frozen audit routes. This builder reads the validated V7
manuscript, V7 submission metadata, V7 Supporting Information and the final
validated V7 figures. The anonymous main file receives continuous line numbers
and figures immediately after the corresponding Results-section lead paragraph.
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph

import build_chapter2_jeb_docx_v1 as legacy

ROOT = Path(__file__).resolve().parents[1]
CH = ROOT / "docs" / "chapter2"
OUT = CH / "submission_package_v7"

FIGURE_SPECS = [
    {
        "number": 1,
        "section": "Multiple capitulum configurations occur within the dominant young radiation",
        "filename": "figure1_v7_diversity_context.png",
        "caption": "Figure 1. Diversity to be assembled within the dominant radiation. Thirty-six of 38 sampled Japanese concepts occur in the dominant radiation, while the authority-backed trait matrix retains multiple capitulum configurations. Missing and conflicting states remain explicit. The scaffold provenance is descriptive context rather than a dated tree.",
        "alt": "Three-panel figure showing 36 of 38 Japanese concepts in the dominant radiation, accepted common-locus scaffold provenance, an authority-backed orientation phyllary and stickiness matrix with missing states, and four dominant-radiation orientation by stickiness configurations.",
    },
    {
        "number": 2,
        "section": "Trait histories occupy unequal evolutionary depths",
        "filename": "figure2_v7_mosaic_depth.png",
        "caption": "Figure 2. Repeated mosaic assembly at unequal evolutionary depths. The figure combines minimum-change burden, exact topology-only relative-depth envelopes, paired-topology ordering, coverage-matched sensitivity and the zero-of-three shared-transition-localization result. Minimum changes are lower bounds; relative depth is not time; topology fractions are sensitivity descriptors.",
        "alt": "Five-panel figure showing minimum changes for orientation phyllary and stickiness, relative-depth envelopes, paired topology ordering of 1000 of 1000, 993 of 1000 and 905 of 1000, coverage-matched central versus strict-tail results, and zero of three trait pairs passing robust shared localization.",
    },
    {
        "number": 3,
        "section": "Orientation transitions track a fixed East-Asian present-niche regime",
        "filename": "figure3_v7_transition_regime_ecology.png",
        "caption": "Figure 3. Orientation ecology is scale-, history- and transition-conditioned. Cross-scale ecological estimands are kept separate; the fixed U-to-D BIO15-up plus BIO1-down transition-regime test is exceptional under exhaustive state-map ranks; the declared falsification ladder bounds the result; and history-conditioned counterfactuals show why the tip-level BIO15 contrast is not ancestry independent.",
        "alt": "Four-panel figure showing cross-scale BIO12 BIO15 and BIO1 evidence, exact finite-map ranks for the fixed orientation transition regime, strict falsification results including the Japan-only boundary, and the counterfactual conditioning ladder from state frequency to recurrence and relative depth.",
    },
    {
        "number": 4,
        "section": "The present transition regime is not supported as the bounded origin regime",
        "filename": "figure4_v7_bounded_history.png",
        "caption": "Figure 4. Bounded orientation history separates descriptive tendency from origin-regime persistence. The sole calendarized U-to-D event spans 94 chronology pairs and four palaeolocation regions. Southern Japan leads descriptively without crossing the 75% dominance gate, and the fixed current BIO15-positive plus BIO1-negative regime matches only 99 of 376 historical scenarios.",
        "alt": "Four-panel figure showing the 0.79 to 0.74 Ma central chronology inside 94 chronology pairs, regional rank-one counts and a 75 percent gate, central BIOCLIM directions, and historical match fractions totaling 99 of 376 scenarios.",
    },
    {
        "number": 5,
        "section": "Coarse historical climate and sea-level regimes do not provide a recurring rescue",
        "filename": "figure5_v7_identifiability_ceiling.png",
        "caption": "Figure 5. Historical identifiability declines from phenotypic assembly to recurring historical cause. Configuration diversity, repeated minimum histories and relative depth are resolved more strongly than event-linked history; the broader diagnostics recover zero of 324 robust climate event classes and zero of 21 robust global sea-level event-metric classes.",
        "alt": "Four-panel identifiability figure showing a seven-step evidence hierarchy, zero of 324 robust climate event classes, zero of 21 robust sea-level event classes, and the conclusion that phenotypic assembly is identifiable farther than a recurring coarse historical cause.",
    },
]


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser()
    p.add_argument("--output-dir", type=Path, default=OUT)
    p.add_argument("--figure-dir", type=Path, required=True)
    return p.parse_args()


def blank_identifying_metadata(doc: Document) -> None:
    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.comments = ""
    props.category = ""
    props.subject = ""
    props.keywords = ""


def assert_anonymous_visible_text(doc: Document) -> None:
    visible = "\n".join(p.text for p in doc.paragraphs).lower()
    forbidden = (
        "github.com/zuizui0223",
        "zuizui0223",
        "corresponding author name",
        "[email]",
    )
    bad = [x for x in forbidden if x in visible]
    if bad:
        raise RuntimeError(f"Identifying or placeholder text remains in anonymous main: {bad}")


def paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def add_figure_after(paragraph: Paragraph, image_path: Path, caption: str, alt: str) -> None:
    if not image_path.exists() or image_path.stat().st_size < 3000:
        raise RuntimeError(f"Missing or unexpectedly small figure: {image_path}")
    image_par = paragraph_after(paragraph)
    image_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = image_par.add_run().add_picture(str(image_path), width=Inches(6.25))
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", caption.split(".", 1)[0])
    caption_par = paragraph_after(image_par)
    try:
        caption_par.style = "Caption"
    except KeyError:
        pass
    caption_par.add_run(caption)


def first_body_paragraph_after_heading(doc: Document, heading_text: str) -> Paragraph:
    pars = list(doc.paragraphs)
    hits = [i for i, p in enumerate(pars) if p.text.strip() == heading_text]
    if len(hits) != 1:
        raise RuntimeError(f"Expected one Results heading {heading_text!r}, found {len(hits)}")
    for p in pars[hits[0] + 1 :]:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name.startswith("Heading"):
            raise RuntimeError(f"No body paragraph found below heading {heading_text!r}")
        return p
    raise RuntimeError(f"No body paragraph found below heading {heading_text!r}")


def embed_main_figures(doc: Document, figure_dir: Path) -> None:
    for spec in FIGURE_SPECS:
        anchor = first_body_paragraph_after_heading(doc, spec["section"])
        add_figure_after(anchor, figure_dir / spec["filename"], spec["caption"], spec["alt"])


def build_main(output_dir: Path, figure_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="Mosaic capitulum assembly through time", line_numbers=True)
    blank_identifying_metadata(doc)
    legacy.render_markdown(
        doc,
        CH / "MANUSCRIPT_JEB_V7_WORKING.md",
        skip_metadata_prefixes=(
            "**Target journal:**",
            "**Manuscript status:**",
            "**Running title:**",
        ),
    )
    assert_anonymous_visible_text(doc)
    embed_main_figures(doc, figure_dir)
    path = output_dir / "Chapter2_JEB_Anonymous_Manuscript_V7.docx"
    legacy.save_document(doc, path)
    return path


def build_title_page(output_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="", line_numbers=False)
    legacy.render_markdown(doc, CH / "JEB_V7_TITLE_PAGE_WORKING.md")
    path = output_dir / "Chapter2_JEB_Title_Page_V7.docx"
    legacy.save_document(doc, path)
    return path


def build_supporting(output_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="Supporting Information", line_numbers=True)
    legacy.render_markdown(doc, CH / "JEB_SUPPORTING_INFORMATION_V3.md")
    path = output_dir / "Chapter2_JEB_Supporting_Information_V3.docx"
    legacy.save_document(doc, path)
    return path


def build_cover_letter(output_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="", line_numbers=False)
    legacy.render_markdown(doc, CH / "JEB_COVER_LETTER_TEMPLATE_V3.md")
    path = output_dir / "Chapter2_JEB_Cover_Letter_V7.docx"
    legacy.save_document(doc, path)
    return path


def main() -> int:
    args = parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    outputs = [
        build_main(args.output_dir, args.figure_dir),
        build_title_page(args.output_dir),
        build_supporting(args.output_dir),
        build_cover_letter(args.output_dir),
    ]
    for path in outputs:
        if not path.exists() or path.stat().st_size < 1000:
            raise RuntimeError(f"Document build failed or unexpectedly small: {path}")
        print(path, path.stat().st_size)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
