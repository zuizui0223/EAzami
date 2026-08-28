#!/usr/bin/env python3
"""Build the double-anonymous JEB DOCX package from reviewed Markdown sources."""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CH = ROOT / "docs" / "chapter2"
OUT = CH / "submission_package"

BODY_FONT = "Times New Roman"
INK = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_FILL = "F2F4F7"


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser()
    p.add_argument("--output-dir", type=Path, default=OUT)
    return p.parse_args()


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = tcMar.find(qn(f"w:{key}"))
        if tag is None:
            tag = OxmlElement(f"w:{key}")
            tcMar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        trPr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def set_run_font(run, size=12, color=INK, bold=None, italic=None, name=BODY_FONT) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_line_numbering(section) -> None:
    sectPr = section._sectPr
    existing = sectPr.find(qn("w:lnNumType"))
    if existing is not None:
        sectPr.remove(existing)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "continuous")
    ln.set(qn("w:distance"), "360")
    sectPr.append(ln)


def configure_document(doc: Document, *, running_header: str, line_numbers: bool) -> None:
    # The stock python-docx template can inherit an odd/even header setting on
    # Windows.  Disable it explicitly so the anonymous running head is present
    # on every page and reviewer line references remain continuous.
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    if line_numbers:
        add_line_numbering(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    heading_specs = {
        "Heading 1": (14, 12, 6),
        "Heading 2": (12, 10, 4),
        "Heading 3": (12, 8, 3),
    }
    for name, (size, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = INK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0

    caption = styles["Caption"]
    caption.font.name = BODY_FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    caption.font.size = Pt(10)
    caption.font.color.rgb = INK
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.0

    if running_header:
        header = section.header
        hp = header.paragraphs[0]
        hp.text = running_header
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        for run in hp.runs:
            set_run_font(run, size=8.5, color=MUTED, italic=True)
    add_page_field(section.footer.paragraphs[0])

    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.comments = "Double-anonymous JEB submission package"
    props.keywords = "Cirsium; phenotypic integration; phylogenetic history"


def add_inline_runs(paragraph, text: str, *, size=12) -> None:
    token_re = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, italic=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, name="Courier New")
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size)


def add_title(doc: Document, text: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=18, bold=True)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(12)
        p2.paragraph_format.line_spacing = 1.0
        run = p2.add_run(subtitle)
        set_run_font(run, size=11, color=MUTED, italic=True)


def add_picture(doc: Document, image_path: Path, alt_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    shape = run.add_picture(str(image_path), width=Inches(6.3))
    docPr = shape._inline.docPr
    docPr.set("descr", alt_text)
    docPr.set("title", image_path.name)
    p.paragraph_format.keep_with_next = True


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [x.strip() for x in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", x) for x in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, matrix: list[list[str]]) -> None:
    if not matrix:
        return
    ncols = max(len(r) for r in matrix)
    for r in matrix:
        r.extend([""] * (ncols - len(r)))
    weights = []
    for col in range(ncols):
        weights.append(max(8, min(50, max(len(re.sub(r"[*`]", "", row[col])) for row in matrix))))
    # Reserve a readable floor for every column, then distribute the remaining
    # width by content.  Applying a floor after proportional allocation can
    # overrun the table and collapse the last column when a table has 7-8 cols.
    min_width = 800 if ncols >= 7 else 900
    extra = 9360 - min_width * ncols
    if extra < 0:
        raise ValueError(f"table has too many columns for readable layout: {ncols}")
    widths = [min_width + int(extra * w / sum(weights)) for w in weights]
    diff = 9360 - sum(widths)
    widths[-1] += diff
    table = doc.add_table(rows=len(matrix), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(matrix):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline_runs(p, value, size=8.5)
            if i == 0:
                for run in p.runs:
                    run.bold = True
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), LIGHT_FILL)
                cell._tc.get_or_add_tcPr().append(shd)
    trPr = table.rows[0]._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trPr.append(header)
    # Do not leave a repeating header as the final row on a page.  Keeping the
    # header paragraphs with the next row lets Word move a short table cleanly.
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = True
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def render_markdown(
    doc: Document,
    source: Path,
    *,
    title_override: str | None = None,
    skip_metadata_prefixes: tuple[str, ...] = (),
    stop_heading: str | None = None,
) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    in_code = False
    idx = 0
    first_title_done = False
    while idx < len(lines):
        raw = lines[idx]
        line = raw.strip()
        if line.startswith("```"):
            in_code = not in_code
            idx += 1
            continue
        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(raw)
            set_run_font(run, size=9, name="Courier New")
            idx += 1
            continue
        if not line or line == "---":
            idx += 1
            continue
        if any(line.startswith(prefix) for prefix in skip_metadata_prefixes):
            idx += 1
            continue
        if line == "<PAGEBREAK>":
            doc.add_page_break()
            idx += 1
            continue
        if line.startswith("# "):
            heading = line[2:].strip()
            if stop_heading and heading == stop_heading:
                break
            if not first_title_done:
                add_title(doc, title_override or heading)
                first_title_done = True
            else:
                doc.add_heading(heading, level=1)
            idx += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            idx += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            idx += 1
            continue
        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            alt, rel = image_match.groups()
            image_path = (source.parent / rel).resolve()
            add_picture(doc, image_path, alt)
            idx += 1
            continue
        if line.startswith("|") and "|" in line[1:]:
            block = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                block.append(lines[idx].strip())
                idx += 1
            add_table(doc, parse_table(block))
            continue
        if re.match(r"^- \[[ xX]\] ", line):
            checked = line[3].lower() == "x"
            text = re.sub(r"^- \[[ xX]\] ", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, ("Completed: " if checked else "Pending: ") + text)
            idx += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            add_inline_runs(p, line[2:])
            idx += 1
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            add_inline_runs(p, re.sub(r"^\d+\. ", "", line))
            idx += 1
            continue
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.right_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            add_inline_runs(p, line[2:])
            for run in p.runs:
                run.italic = True
            idx += 1
            continue
        is_figure_caption = line.startswith("**Figure ")
        is_table_caption = line.startswith("**Table ")
        is_caption = is_figure_caption or is_table_caption
        p = doc.add_paragraph(style="Caption" if is_caption else "Normal")
        if is_caption:
            p.paragraph_format.keep_with_next = is_table_caption
        add_inline_runs(p, line, size=10 if is_caption else 12)
        idx += 1


def save_document(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_main(output_dir: Path) -> Path:
    doc = Document()
    configure_document(doc, running_header="", line_numbers=True)
    render_markdown(
        doc,
        CH / "MANUSCRIPT_JEB_V3.md",
        skip_metadata_prefixes=("**Target journal:**", "**Manuscript status:**", "**Running title:**", "**Word-limit contract:**"),
        stop_heading="Submission completion gates",
    )
    path = output_dir / "Chapter2_JEB_Anonymous_Manuscript_V3.docx"
    save_document(doc, path)
    return path


def build_title_page(output_dir: Path) -> Path:
    doc = Document()
    configure_document(doc, running_header="", line_numbers=False)
    render_markdown(doc, CH / "JEB_TITLE_PAGE_TEMPLATE_V1.md")
    path = output_dir / "Chapter2_JEB_Title_Page_TEMPLATE_V1.docx"
    save_document(doc, path)
    return path


def build_supporting(output_dir: Path) -> Path:
    doc = Document()
    configure_document(doc, running_header="", line_numbers=True)
    render_markdown(doc, CH / "JEB_SUPPORTING_INFORMATION_V1.md")
    path = output_dir / "Chapter2_JEB_Supporting_Information_V1.docx"
    save_document(doc, path)
    return path


def build_cover_letter(output_dir: Path) -> Path:
    doc = Document()
    configure_document(doc, running_header="", line_numbers=False)
    render_markdown(doc, CH / "JEB_COVER_LETTER_TEMPLATE_V1.md")
    path = output_dir / "Chapter2_JEB_Cover_Letter_TEMPLATE_V1.docx"
    save_document(doc, path)
    return path


def main() -> int:
    args = parse_args()
    outputs = [
        build_main(args.output_dir),
        build_title_page(args.output_dir),
        build_supporting(args.output_dir),
        build_cover_letter(args.output_dir),
    ]
    for path in outputs:
        print(path.relative_to(ROOT))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
