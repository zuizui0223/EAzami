#!/usr/bin/env python3
"""Build the active double-anonymous JEB package from V4/V2 Markdown sources.

This intentionally reuses the rendering/layout helpers from the frozen legacy
v1 builder while switching only the active source and output version routing.
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document

import build_chapter2_jeb_docx_v1 as legacy

ROOT = Path(__file__).resolve().parents[1]
CH = ROOT / "docs" / "chapter2"
OUT = CH / "submission_package"


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser()
    p.add_argument("--output-dir", type=Path, default=OUT)
    return p.parse_args()


def build_main(output_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="", line_numbers=True)
    legacy.render_markdown(
        doc,
        CH / "MANUSCRIPT_JEB_V4.md",
        skip_metadata_prefixes=(
            "**Target journal:**",
            "**Manuscript status:**",
            "**Running title:**",
            "**Word-limit contract:**",
        ),
        stop_heading="Submission completion gates",
    )
    path = output_dir / "Chapter2_JEB_Anonymous_Manuscript_V4.docx"
    legacy.save_document(doc, path)
    return path


def build_title_page(output_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="", line_numbers=False)
    legacy.render_markdown(doc, CH / "JEB_TITLE_PAGE_TEMPLATE_V1.md")
    path = output_dir / "Chapter2_JEB_Title_Page_TEMPLATE_V1.docx"
    legacy.save_document(doc, path)
    return path


def build_supporting(output_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="", line_numbers=True)
    legacy.render_markdown(doc, CH / "JEB_SUPPORTING_INFORMATION_V2.md")
    path = output_dir / "Chapter2_JEB_Supporting_Information_V2.docx"
    legacy.save_document(doc, path)
    return path


def build_cover_letter(output_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="", line_numbers=False)
    legacy.render_markdown(doc, CH / "JEB_COVER_LETTER_TEMPLATE_V1.md")
    path = output_dir / "Chapter2_JEB_Cover_Letter_TEMPLATE_V1.docx"
    legacy.save_document(doc, path)
    return path


def display_path(path: Path) -> str:
    try:
        return str(path.relative_to(ROOT))
    except ValueError:
        return str(path)


def main() -> int:
    args = parse_args()
    outputs = [
        build_main(args.output_dir),
        build_title_page(args.output_dir),
        build_supporting(args.output_dir),
        build_cover_letter(args.output_dir),
    ]
    for path in outputs:
        print(display_path(path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
