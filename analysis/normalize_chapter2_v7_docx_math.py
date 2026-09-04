#!/usr/bin/env python3
"""Normalize the small amount of Markdown/LaTeX math in the rendered V7 DOCX.

The validated scientific Markdown remains unchanged. This presentation-only
postprocessor removes display delimiters and replaces the two simple formulas
with plain Unicode/Word-readable notation before the submission artifact is
frozen.
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def remove_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def replace_text(paragraph, text: str, *, centered: bool = False) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    if centered:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def normalize(path: Path) -> None:
    doc = Document(path)
    for paragraph in list(doc.paragraphs):
        text = paragraph.text
        stripped = text.strip()
        if stripped in {r"\[", r"\]"}:
            remove_paragraph(paragraph)
            continue
        if stripped == r"D = \frac{N-d}{N-1}.":
            replace_text(paragraph, "D = (N − d)/(N − 1).", centered=True)
            continue
        new = text.replace(r"\(D=1\)", "D = 1")
        new = new.replace(r"\(\binom{9}{4}=126\)", "C(9, 4) = 126")
        if new != text:
            replace_text(paragraph, new)
    doc.save(path)

    check = Document(path)
    visible = "\n".join(p.text for p in check.paragraphs)
    forbidden = (r"\[", r"\]", r"\frac", r"\binom", r"\(", r"\)")
    bad = [token for token in forbidden if token in visible]
    if bad:
        raise RuntimeError(f"raw LaTeX remains in V7 DOCX: {bad}")
    if "D = (N − d)/(N − 1)." not in visible or "C(9, 4) = 126" not in visible:
        raise RuntimeError("normalized V7 formulas missing after DOCX rewrite")
    print(f"normalized V7 DOCX math: {path}")


def main() -> int:
    p = argparse.ArgumentParser()
    p.add_argument("docx", type=Path)
    args = p.parse_args()
    normalize(args.docx)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
