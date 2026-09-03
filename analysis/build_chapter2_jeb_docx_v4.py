#!/usr/bin/env python3
"""Build the active double-anonymous JEB V6 package.

V5 and earlier builders remain frozen audit routes. This builder reads only the
V6 final manuscript, V6 figures, V4 Supporting Information, and V3 submission
metadata templates.
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph

import build_chapter2_jeb_docx_v1 as legacy

ROOT = Path(__file__).resolve().parents[1]
CH = ROOT / "docs" / "chapter2"
OUT = CH / "submission_package_v6"

FIGURE_SPECS = [
    {
        "number": 1,
        "marker": "(Fig. 1",
        "filename": "figure1_recurrence_depth.png",
        "caption": "Figure 1. Capitulum modules repeatedly differentiated at unequal evolutionary depths. Trait-state coverage, minimum unordered changes, bootstrap-median relative lineage-depth envelopes and the zero-of-three shared-localization result are shown. Relative lineage depth is topology-only and is not calendar time or an evolutionary rate.",
        "alt": "Four-panel figure for orientation, phyllary posture and stickiness. Resolved concept counts are 20, 10 and 13; minimum changes are 4 to 6, 3 and 5; relative-depth envelopes are 0.795 to 0.994, 0.695 to 1.000 and 0.937 to 0.954; zero of three trait pairs passes robust shared localization.",
    },
    {
        "number": 2,
        "marker": "(Fig. 2",
        "filename": "figure2_calendar_identifiability.png",
        "caption": "Figure 2. Calendar-time identifiability is the main historical bottleneck. The evidence funnel shows that only one current capitulum transition reaches the full trait-transition, calendar, palaeolocation and historical-environment gate, despite well-resolved repeated histories for several modules.",
        "alt": "Three-panel figure showing a historical evidence funnel across orientation, phyllary posture, stickiness and colour; counts of full-gate and weaker dated evidence classes; and the conclusion that repeated trait history is much better identified than repeated historical cause.",
    },
    {
        "number": 3,
        "marker": "(Fig. 3",
        "filename": "figure3_orientation_uncertainty.png",
        "caption": "Figure 3. A clean central-date orientation story disappears under historical uncertainty. The core-Nipponocirsium orientation event is evaluated across 94 admissible chronologies and four palaeolocation regions. The central 0.79–0.74 Ma pair has coherent climate directions, but no BIO1, BIO4, BIO12 or BIO15 signed direction, environmental level, absolute change or temporal variability passes the full robust gate; no global sea-level metric survives all chronologies.",
        "alt": "Four-panel orientation figure showing the bounded branch, central-date climate directions, a matrix with zero robust climate variables for direction, level, absolute change and variability, and BIO1 sub-threshold matched-window tendencies. De Boer sea level covers all 94 chronologies but yields no robust metric.",
    },
    {
        "number": 4,
        "marker": "(Fig. 4",
        "filename": "figure4_repeated_trigger_ceiling.png",
        "caption": "Figure 4. Broader dated lineage contexts do not recover one recurring coarse trigger. A 17-BIOCLIM atlas across six dated lineage contexts yields zero of 324 robust event-level classes, and three representative clades across seven global sea-level metrics yield zero of 21 robust event-metric classes.",
        "alt": "Four-panel figure summarizing 17 BIOCLIM variables, six dated contexts and three clade groups, with zero of 324 robust climate event classes and zero of 21 robust sea-level event classes. Local fragmentation, connectivity, biotic interactions and lineage-specific exposure remain open.",
    },
    {
        "number": 5,
        "marker": "(Fig. 5",
        "filename": "figure5_public_data_endpoint.png",
        "caption": "Figure 5. History is substantially better resolved than historical cause. The evidence ladder progresses from strong recurrence and relative-depth results through sparse calendar event placement to the final outcome that no recurring tested coarse historical trigger is identified under current public-data uncertainty.",
        "alt": "Two-panel synthesis figure. An evidence ladder shows strong recurrence and relative depth, zero of three robust shared-history pairs, sparse calendar event placement and an unidentified recurring historical trigger. The final classification states that repeated differentiation is resolved but the recurring tested historical trigger is not identified.",
    },
]


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser()
    p.add_argument("--output-dir", type=Path, default=OUT)
    p.add_argument("--figure-dir", type=Path, default=CH / "figures_v6")
    return p.parse_args()


def blank_identifying_metadata(doc: Document) -> None:
    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.comments = ""
    props.category = ""
    props.subject = ""
    props.keywords = ""


def paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def add_figure_after(paragraph: Paragraph, image_path: Path, caption: str, alt: str) -> None:
    if not image_path.exists() or image_path.stat().st_size < 3000:
        raise RuntimeError(f"Missing or unexpectedly small figure: {image_path}")
    image_par = paragraph_after(paragraph)
    image_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = image_par.add_run().add_picture(str(image_path), width=Inches(6.35))
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", caption.split(".", 1)[0])
    caption_par = paragraph_after(image_par)
    try:
        caption_par.style = "Caption"
    except KeyError:
        pass
    caption_par.add_run(caption)


def embed_main_figures(doc: Document, figure_dir: Path) -> None:
    original = list(doc.paragraphs)
    used: set[int] = set()
    for spec in FIGURE_SPECS:
        matches = [p for p in original if spec["marker"] in p.text]
        if not matches:
            raise RuntimeError(f"No first-mention paragraph found for Figure {spec['number']}")
        add_figure_after(matches[0], figure_dir / spec["filename"], spec["caption"], spec["alt"])
        used.add(spec["number"])
    if used != {1, 2, 3, 4, 5}:
        raise RuntimeError(f"Incomplete figure insertion: {sorted(used)}")


def build_main(output_dir: Path, figure_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="", line_numbers=True)
    blank_identifying_metadata(doc)
    legacy.render_markdown(
        doc,
        CH / "MANUSCRIPT_JEB_V6_FINAL.md",
        skip_metadata_prefixes=(
            "**Target journal:**",
            "**Manuscript status:**",
            "**Running title:**",
        ),
        stop_heading="Submission completion gates",
    )
    embed_main_figures(doc, figure_dir)
    path = output_dir / "Chapter2_JEB_Anonymous_Manuscript_V6.docx"
    legacy.save_document(doc, path)
    return path


def build_title_page(output_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="", line_numbers=False)
    legacy.render_markdown(doc, CH / "JEB_TITLE_PAGE_TEMPLATE_V3.md")
    path = output_dir / "Chapter2_JEB_Title_Page_TEMPLATE_V3.docx"
    legacy.save_document(doc, path)
    return path


def build_supporting(output_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="", line_numbers=True)
    legacy.render_markdown(doc, CH / "JEB_SUPPORTING_INFORMATION_V4.md")
    path = output_dir / "Chapter2_JEB_Supporting_Information_V4.docx"
    legacy.save_document(doc, path)
    return path


def build_cover_letter(output_dir: Path) -> Path:
    doc = Document()
    legacy.configure_document(doc, running_header="", line_numbers=False)
    legacy.render_markdown(doc, CH / "JEB_COVER_LETTER_TEMPLATE_V3.md")
    path = output_dir / "Chapter2_JEB_Cover_Letter_TEMPLATE_V3.docx"
    legacy.save_document(doc, path)
    return path


def display_path(path: Path) -> str:
    try:
        return str(path.relative_to(ROOT))
    except ValueError:
        return str(path)


def main() -> int:
    args = parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    outputs = [
        build_main(args.output_dir, args.figure_dir),
        build_title_page(args.output_dir),
        build_supporting(args.output_dir),
        build_cover_letter(args.output_dir),
    ]
    for path in outputs:
        if not path.exists() or path.stat().st_size < 1000:
            raise RuntimeError(f"Document build failed or unexpectedly small: {path}")
        print(display_path(path), path.stat().st_size)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
